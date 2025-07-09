import os
from dotenv import load_dotenv
import streamlit as st
import sqlite3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn as qn_ns
import tempfile
import re
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
import datetime
from openai import AzureOpenAI

# 환경 변수 로드
load_dotenv()

# 데이터 캐시 비우기
st.cache_data.clear()
st.cache_resource.clear()

# Azure OpenAI 설정
openai_endpoint = os.getenv("OPENAI_ENDPOINT")
openai_api_key = os.getenv("OPENAI_KEY")
chat_model = os.getenv("CHAT_MODEL")
# 환경변수 설정
AZURE_STORAGE_CONNECTION_STRING = os.getenv("STORAGE_CONN_STR")

STORAGE_ACCOUNT_NAME = os.getenv("STORAGE_ACCOUNT_NAME")
WORD_CONTAINER_NAME = os.getenv("WORD_CONTAINER_NAME")
EML_DB_NAME = os.getenv("EML_DB_NAME")


def get_all_eml_records():
    """모든 EML 레코드 조회"""
    try:
        conn = sqlite3.connect(EML_DB_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, original_filename, subject, body_text
            FROM eml_reports 
            WHERE body_text IS NOT NULL AND body_text != ''
            ORDER BY upload_time DESC
        ''')
        
        records = cursor.fetchall()
        conn.close()
        
        return records
    except Exception as e:
        st.error(f"데이터베이스 조회 실패: {str(e)}")
        return []

def read_sample_document(file_path: str) -> str:
    """샘플 문서의 내용을 읽어서 텍스트로 반환"""
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 빈 문단 제외
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"샘플 문서를 읽을 수 없습니다: {str(e)}"

def generate_llm_report(body_text: str, sample_format: str) -> str:
    """Azure OpenAI를 사용하여 LLM 보고서 생성"""
    
    try:
        if not all([openai_endpoint, openai_api_key, chat_model]):
            return "Azure OpenAI 설정이 완료되지 않았습니다. 환경변수를 확인해주세요."
        
        # Azure OpenAI 클라이언트 초기화
        client = AzureOpenAI(
            azure_endpoint=openai_endpoint,
            api_key=openai_api_key,
            api_version="2024-02-15-preview"
        )
        
        # 프롬프트 생성
        prompt = f"""
다음 샘플 장애보고서의 형식을 참고하여, 주어진 본문 내용을 바탕으로 전문적인 장애보고서를 작성해주세요.

샘플 장애보고서 형식:
{sample_format}

본문 내용:
{body_text}

위의 본문 내용을 분석하여 샘플 형식에 맞는 전문적인 장애보고서를 작성해주세요. 
- 장애 발생 시간, 원인, 영향도, 조치사항 등을 명확히 구분하여 작성
- 기술적인 내용은 정확하고 이해하기 쉽게 설명
- 보고서 형식은 샘플과 동일하게 유지
- 표 형식의 정보는 다음과 같이 마크다운 표 형식으로 작성:
  | 항목 | 내용 |
  |------|------|
  | 장애발생일시 | 2024-XX-XX XX:XX |
  | 장애해결일시 | 2024-XX-XX XX:XX |
  | 영향도 | 상/중/하 |
  | 장애원인 | 구체적 원인 |
  | 조치사항 | 구체적 조치 내용 |
"""
        
        # OpenAI API 호출
        response = client.chat.completions.create(
            model=chat_model,
            messages=[
                {"role": "system", "content": "당신은 IT 장애보고서 작성 전문가입니다. 주어진 정보를 바탕으로 정확하고 전문적인 장애보고서를 작성해주세요. 표 형식의 정보는 마크다운 표 형식으로 작성해주세요."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        return f"LLM 보고서 생성 중 오류가 발생했습니다: {str(e)}"

def parse_markdown_table(table_text):
    """마크다운 표 텍스트를 파싱하여 테이블 데이터 반환"""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:  # 최소 헤더, 구분자, 데이터 1행
        return None
    
    # 헤더 행 파싱
    header_line = lines[0].strip()
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]  # 첫 번째와 마지막 빈 요소 제거
    
    # 구분자 행 건너뛰기 (lines[1])
    
    # 데이터 행들 파싱
    data_rows = []
    for i in range(2, len(lines)):
        line = lines[i].strip()
        if line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # 첫 번째와 마지막 빈 요소 제거
            if len(cells) == len(headers):
                data_rows.append(cells)
    
    return headers, data_rows

def create_word_table(doc, headers, data_rows):
    """Word 문서에 표 추가"""
    # 표 생성 (헤더 포함)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 표 너비 설정
    table.autofit = False
    table.allow_autofit = False
    
    # 헤더 행 설정
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        # 헤더 셀 스타일링
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        
        # 헤더 배경색 설정 (회색)
        shading = OxmlElement('w:shd')
        shading.set(qn_ns('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # 데이터 행 추가
    for row_data in data_rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            cell.text = cell_data
            # 데이터 셀 스타일링
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
            run.font.name = '맑은 고딕'
            run.font.size = Pt(11)
    
    # 표 열 너비 자동 조정
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(2.5)  # 기본 너비 설정
    
    return table

def generate_word_from_llm_report(llm_report: str, filename: str) -> str:
    """LLM 생성 보고서를 기반으로 Word 문서 생성 (마크다운을 Word 스타일로 변환)"""
    # 새로운 Word 문서 생성
    doc = Document()
    
    # 문서 제목 추가
    title = doc.add_heading('장애보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # LLM 보고서 내용을 줄별로 분리하여 처리
    lines = llm_report.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # 빈 줄은 빈 문단으로 추가
            doc.add_paragraph()
            i += 1
            continue
        
        # 마크다운 표 감지
        if line.startswith('|') and '|' in line:
            # 표 시작 감지
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip() and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1
            
            if len(table_lines) >= 3:  # 최소 헤더, 구분자, 데이터 1행
                # 표 파싱 및 생성
                table_text = '\n'.join(table_lines)
                table_data = parse_markdown_table(table_text)
                
                if table_data:
                    headers, data_rows = table_data
                    create_word_table(doc, headers, data_rows)
                    doc.add_paragraph()  # 표 다음에 빈 줄 추가
                    i = j
                    continue
                else:
                    # 표 파싱 실패 시 일반 텍스트로 처리
                    paragraph = doc.add_paragraph()
                    add_formatted_text(paragraph, line)
                    set_korean_font(paragraph)
                    i += 1
                    continue
            else:
                # 표가 아닌 일반 텍스트로 처리
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, line)
                set_korean_font(paragraph)
                i += 1
                continue
        
        # 마크다운 헤딩 처리
        if line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            set_korean_font(heading)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            set_korean_font(heading)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            set_korean_font(heading)
        elif line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            set_korean_font(heading)
        # 순서 있는 리스트 처리 (숫자로 시작하는 패턴)
        elif re.match(r'^\d+\.\s+', line):
            # 숫자 제거하고 내용만 추출
            content = re.sub(r'^\d+\.\s+', '', line)
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Number'
            add_formatted_text(paragraph, content)
            set_korean_font(paragraph)
        # 순서 없는 리스트 처리
        elif line.startswith('- ') or line.startswith('* '):
            content = line[2:]
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Bullet'
            add_formatted_text(paragraph, content)
            set_korean_font(paragraph)
        # 전체가 볼드인 경우 (제목처럼 사용)
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(12)
            set_korean_font(paragraph)
        # 일반 텍스트 (인라인 포맷팅 포함)
        else:
            paragraph = doc.add_paragraph()
            add_formatted_text(paragraph, line)
            set_korean_font(paragraph)
        
        i += 1
    
    # 문서 전체 스타일 설정
    set_document_style(doc)
    
    # 임시 파일에 저장
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_path)
    return temp_path

def add_formatted_text(paragraph, text):
    """문단에 인라인 포맷팅이 적용된 텍스트 추가"""
    # 볼드 텍스트 패턴 처리
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            # 볼드 텍스트
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # 일반 텍스트
            run = paragraph.add_run(part)

def set_korean_font(paragraph):
    """한글 폰트 설정"""
    for run in paragraph.runs:
        run.font.name = '맑은 고딕'
        run.font.size = Pt(11)
        
        # 한글 폰트 설정을 위한 XML 처리
        r = run._element
        rPr = r.get_or_add_rPr()
        
        # 동아시아 폰트 설정
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:hint'), 'eastAsia')

def set_document_style(doc):
    """문서 전체 스타일 설정"""
    # 기본 스타일 설정
    styles = doc.styles
    
    # Normal 스타일 설정
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '맑은 고딕'
    normal_font.size = Pt(11)
    
    # 문단 간격 설정
    paragraph_format = normal_style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    
    # 헤딩 스타일 설정
    for level in range(1, 5):
        heading_style = styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = '맑은 고딕'
        heading_font.bold = True
        
        # 헤딩 레벨별 크기 설정
        if level == 1:
            heading_font.size = Pt(16)
        elif level == 2:
            heading_font.size = Pt(14)
        elif level == 3:
            heading_font.size = Pt(13)
        else:  # level 4
            heading_font.size = Pt(12)
        
        # 헤딩 간격 설정
        heading_format = heading_style.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
    
    # 리스트 스타일 설정
    try:
        list_bullet_style = styles['List Bullet']
        list_bullet_font = list_bullet_style.font
        list_bullet_font.name = '맑은 고딕'
        list_bullet_font.size = Pt(11)
        
        list_number_style = styles['List Number']
        list_number_font = list_number_style.font
        list_number_font.name = '맑은 고딕'
        list_number_font.size = Pt(11)
    except KeyError:
        # 스타일이 없는 경우 무시
        pass

def upload_to_azure_word_blob(file_path, filename):
    """Azure Blob Storage에 파일 업로드하고 다운로드 URL 반환"""
    try:
        if not AZURE_STORAGE_CONNECTION_STRING:
            return False, None, "Azure Storage 연결 문자열이 설정되지 않았습니다."
        
        # 연결 문자열 유효성 검사
        is_valid, message = validate_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        if not is_valid:
            return False, None, f"연결 문자열 오류: {message}"
        
        # Blob Service Client 생성
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        
        # 타임스탬프를 포함한 고유한 파일명 생성
        timestamp = datetime.datetime.now(datetime.UTC).strftime("%Y%m%d_%H%M%S")
        blob_name = f"{timestamp}_{filename}"
        
        # 컨테이너 존재 확인 및 생성
        try:
            container_client = blob_service_client.get_container_client(WORD_CONTAINER_NAME)
            container_client.get_container_properties()
        except Exception as e:
            if "ContainerNotFound" in str(e):
                container_client.create_container()
        
        # Blob 업로드
        blob_client = blob_service_client.get_blob_client(
            container=WORD_CONTAINER_NAME, 
            blob=blob_name
        )
        
        # 파일 업로드
        with open(file_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
        
        # 다운로드 URL 생성 (SAS 토큰 포함)
        # 연결 문자열에서 계정 키 추출
        account_key = None
        for part in AZURE_STORAGE_CONNECTION_STRING.split(';'):
            if part.startswith('AccountKey='):
                account_key = part.split('=', 1)[1]
                break
        
        if account_key:
            # SAS 토큰 생성 (24시간 유효)
            sas_token = generate_blob_sas(
                account_name=STORAGE_ACCOUNT_NAME,
                container_name=WORD_CONTAINER_NAME,
                blob_name=blob_name,
                account_key=account_key,
                permission=BlobSasPermissions(read=True),
                expiry=datetime.datetime.now(datetime.UTC) + datetime.timedelta(hours=24)
            )
            
            # 다운로드 URL 생성
            blob_url = f"https://{STORAGE_ACCOUNT_NAME}.blob.core.windows.net/{WORD_CONTAINER_NAME}/{blob_name}?{sas_token}"
        else:
            # 계정 키를 찾을 수 없으면 기본 URL만 반환
            blob_url = f"https://{STORAGE_ACCOUNT_NAME}.blob.core.windows.net/{WORD_CONTAINER_NAME}/{blob_name}"
        
        return True, blob_url, None
        
    except Exception as e:
        error_message = str(e)
        if "Connection string is either blank or malformed" in error_message:
            error_message = "Azure Storage 연결 문자열이 올바르지 않습니다. 설정을 확인해주세요."
        return False, None, error_message

def validate_connection_string(connection_string):
    """연결 문자열 유효성 검사"""
    if not connection_string:
        return False, "연결 문자열이 비어있습니다."
    
    # 기본적인 Azure Storage 연결 문자열 형식 확인
    required_parts = ['AccountName=', 'AccountKey=']
    missing_parts = [part for part in required_parts if part not in connection_string]
    
    if missing_parts:
        return False, f"연결 문자열에 필수 요소가 누락되었습니다: {', '.join(missing_parts)}"
    
    return True, "유효한 연결 문자열입니다."    

# === Streamlit 앱 ===
st.title("📧 EML 본문 조회")

# 데이터베이스 연결 확인
try:
    conn = sqlite3.connect(EML_DB_NAME)
    conn.close()
except Exception as e:
    st.error(f"❌ 데이터베이스 연결 실패: {str(e)}")
    st.stop()

# 레코드 조회
records = get_all_eml_records()

if not records:
    st.info("조회된 레코드가 없습니다.")
else:
    # 레코드 선택
    record_options = [f"ID {record[0]} - {record[1]}" for record in records]
    selected_index = st.selectbox("조회할 레코드를 선택하세요:", range(len(records)), format_func=lambda x: record_options[x])
    
    # 선택된 레코드의 본문 표시
    if selected_index is not None:
        selected_record = records[selected_index]
        body_text = selected_record[3]
        
        # 본문 내용 표시
        st.subheader("본문 내용")
        with st.expander("본문 내용 보기", expanded=True):
            st.markdown(f"```\n{body_text}\n```")
        
        # 샘플 장애보고서 내용 표시
        st.subheader("샘플 장애보고서")
        sample_content = read_sample_document("data/docx/iap-report-sample1.docx")
        with st.expander("샘플 장애보고서 내용 보기", expanded=True):
            st.markdown(f"```\n{sample_content}\n```")

        # LLM 생성 보고서 섹션 추가
        st.subheader("🤖 LLM 생성 보고서")

        with st.spinner("LLM이 보고서를 생성하고 있습니다..."):
            llm_report = generate_llm_report(body_text, sample_content)
        
        with st.expander("LLM 생성 보고서 보기", expanded=True):
            st.markdown(llm_report)

        st.markdown("---")
        if st.button("📄 Word 파일 생성 및 업로드"):
            with st.spinner("Word 파일을 생성하고 업로드하는 중..."):
                # 수정: st.markdown(llm_report) 대신 llm_report 직접 전달
                word_file_path = generate_word_from_llm_report(llm_report, "generated_report.docx")
                blob_name = f"iap-report-{selected_record[0]}.docx"
                
                success, blob_url, error_msg = upload_to_azure_word_blob(word_file_path, blob_name)
                
                # 임시 파일 삭제
                if os.path.exists(word_file_path):
                    os.unlink(word_file_path)
                
                if success and blob_url:
                    st.success("✅ Word 파일이 생성되어 Azure에 업로드되었습니다.")
                    st.markdown(f"**📥 [파일 다운로드]({blob_url})**")
                    st.info("💡 다운로드 링크는 24시간 동안 유효합니다.")
                else:
                    st.error(f"❌ 업로드 실패: {error_msg}")