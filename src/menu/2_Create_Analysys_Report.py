# 통합 Streamlit 앱: EML 업로드 + 보고서 생성

import os
import sqlite3
import datetime
import re
import tempfile
import email
from email import policy
from dotenv import load_dotenv
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from openai import AzureOpenAI
import base64
import binascii

# 환경변수 로드
load_dotenv()

# 공통 설정 로드
STORAGE_CONN_STR = os.getenv("STORAGE_CONN_STR")
STORAGE_ACCOUNT_NAME = os.getenv("STORAGE_ACCOUNT_NAME")
EML_CONTAINER_NAME = os.getenv("EML_CONTAINER_NAME")
WORD_CONTAINER_NAME = os.getenv("WORD_CONTAINER_NAME")
EML_DB_NAME = os.getenv("EML_DB_NAME")
OPENAI_ENDPOINT = os.getenv("OPENAI_ENDPOINT")
OPENAI_KEY = os.getenv("OPENAI_KEY")
CHAT_MODEL = os.getenv("CHAT_MODEL")

# === 공통 유틸 함수 ===
def validate_connection_string(connection_string):
    if not connection_string:
        return False, "연결 문자열이 비어있습니다."
    required_parts = ['AccountName=', 'AccountKey=']
    missing_parts = [p for p in required_parts if p not in connection_string]
    if missing_parts:
        return False, f"누락 항목: {', '.join(missing_parts)}"
    return True, "정상 연결 문자열"

def init_database():
    try:
        conn = sqlite3.connect(EML_DB_NAME)
        conn.execute('''CREATE TABLE IF NOT EXISTS eml_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            original_filename TEXT, blob_name TEXT,
            subject TEXT, sender TEXT, recipient TEXT,
            date_sent TEXT, body_text TEXT,
            attachments TEXT, file_size INTEGER,
            upload_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
        conn.close()
        return True, "DB 초기화 성공"
    except Exception as e:
        return False, str(e)

def parse_eml(content_str):
    try:
        msg = email.message_from_string(content_str, policy=policy.default)
        body_text, attachments = '', []
        for part in msg.walk():
            if part.get_content_type() == "text/plain" and not part.get("Content-Disposition"):
                body_text = part.get_content()
            elif part.get_filename():
                attachments.append(part.get_filename())
        return {
            'from': msg.get('From', ''), 'to': msg.get('To', ''),
            'subject': msg.get('Subject', ''), 'date': msg.get('Date', ''),
            'body_text': body_text, 'attachments': attachments
        }, None
    except Exception as e:
        return None, str(e)

def upload_blob(content, filename, container_name):
    try:
        client = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
        container = client.get_container_client(container_name)
        try: container.get_container_properties()
        except: container.create_container()
        blob_name = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_") + filename
        blob = container.get_blob_client(blob_name)
        blob.upload_blob(content, overwrite=True)
        return True, blob_name, None
    except Exception as e:
        return False, None, str(e)

def insert_eml(parsed, orig_name, blob_name, size):
    try:
        conn = sqlite3.connect(EML_DB_NAME)
        cursor = conn.cursor()
        cursor.execute('''INSERT INTO eml_reports (original_filename, blob_name, subject, sender, recipient, 
                          date_sent, body_text, attachments, file_size) 
                          VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                       (orig_name, blob_name, parsed['subject'], parsed['from'], parsed['to'], parsed['date'],
                        parsed['body_text'], ', '.join(parsed['attachments']), size))
        conn.commit()
        return True, cursor.lastrowid, None
    except Exception as e:
        return False, None, str(e)

def generate_llm_report(body_text, sample_format):
    try:
        client = AzureOpenAI(azure_endpoint=OPENAI_ENDPOINT, api_key=OPENAI_KEY, api_version="2024-02-15-preview")
        prompt = f"""샘플 형식:
{sample_format}
본문:
{body_text}
샘플을 참고해 전문적 보고서를 작성하세요."""
        res = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": "당신은 IT 보고서 전문가입니다."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        return res.choices[0].message.content
    except Exception as e:
        return str(e)

def generate_word(text, filename):
    doc = Document()
    doc.add_heading('장애보고서', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in text.split('\n'):
        doc.add_paragraph(line)
    for style in doc.styles:  # 간단 스타일
        if style.name == 'Normal':
            style.font.name = '맑은 고딕'
            style.font.size = Pt(11)
    path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(path)
    return path

def upload_word(path, filename):
    client = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
    blob_name = f"{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}"
    blob = client.get_blob_client(WORD_CONTAINER_NAME, blob_name)

    # Word 파일 업로드
    with open(path, 'rb') as f:
        blob.upload_blob(f, overwrite=True)

    # AccountKey 추출 및 base64 유효성 검증
    try:
        key = next(p.split('=', 1)[1] for p in STORAGE_CONN_STR.split(';') if p.startswith('AccountKey='))
        # 패딩 보정 (길이가 4의 배수가 아니면 '=' 추가)
        missing_padding = len(key) % 4
        if missing_padding:
            key += '=' * (4 - missing_padding)
        # base64 디코딩 검증
        base64.b64decode(key)
    except (StopIteration, binascii.Error, base64.binascii.Error) as e:
        raise ValueError(f"❌ AccountKey가 잘못되었거나 base64 형식이 아닙니다: {e}")

    # SAS 토큰 생성
    sas = generate_blob_sas(
        account_name=STORAGE_ACCOUNT_NAME,
        container_name=WORD_CONTAINER_NAME,
        blob_name=blob_name,
        account_key=key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=24)
    )

    return f"https://{STORAGE_ACCOUNT_NAME}.blob.core.windows.net/{WORD_CONTAINER_NAME}/{blob_name}?{sas}"

# === Streamlit 인터페이스 ===
st.set_page_config(page_title="EML 리포트 통합앱", layout="wide")
st.title("📨 장애보고 자동 생성기")

# DB 및 연결 확인
if not all([STORAGE_CONN_STR, STORAGE_ACCOUNT_NAME, EML_CONTAINER_NAME, EML_DB_NAME, OPENAI_ENDPOINT, OPENAI_KEY]):
    st.error("환경변수가 부족합니다. .env 파일을 확인해주세요.")
    st.stop()
init_database()

# 1단계: EML 업로드
st.header("1️⃣ EML 업로드 및 저장")
uploaded = st.file_uploader("EML 파일을 업로드하세요", type='eml')
if uploaded:
    content = uploaded.read().decode('utf-8', errors='ignore')
    parsed, err = parse_eml(content)
    if err:
        st.error(f"파싱 오류: {err}")
    else:
        st.write(f"제목: {parsed['subject']}")
        st.write(f"발신자: {parsed['from']}, 수신자: {parsed['to']}")
        st.text_area("본문 미리보기", parsed['body_text'], height=200)
        if st.button("☁️ 업로드 실행"):
            ok, blob_name, err = upload_blob(content.encode(), uploaded.name, EML_CONTAINER_NAME)
            if ok:
                ok_db, rid, err_db = insert_eml(parsed, uploaded.name, blob_name, len(content))
                st.success(f"업로드 및 저장 성공 (ID: {rid})")
            else:
                st.error(f"업로드 실패: {err}")

# 2단계: 보고서 생성
st.header("2️⃣ 보고서 생성 및 다운로드")
conn = sqlite3.connect(EML_DB_NAME)
rows = conn.execute("SELECT id, subject, body_text FROM eml_reports ORDER BY id DESC").fetchall()
conn.close()

if rows:
    options = [f"ID {r[0]} - {r[1]}" for r in rows]
    idx = st.selectbox("생성할 레코드 선택", list(range(len(rows))), format_func=lambda i: options[i])
    sel = rows[idx]
    st.write("본문:")
    st.code(sel[2])
    sample = Document("data/docx/iap-report-sample1.docx")
    sample_text = '\n'.join(p.text for p in sample.paragraphs if p.text.strip())
    if st.button("🤖 보고서 생성 및 Word 저장"):
        with st.spinner("LLM 보고서 생성 중..."):
            llm_report = generate_llm_report(sel[2], sample_text)
            word_path = generate_word(llm_report, f"report_{sel[0]}.docx")
            url = upload_word(word_path, f"report_{sel[0]}.docx")
            st.success("보고서 업로드 완료!")
            st.markdown(f"📥 [다운로드 링크]({url})")
