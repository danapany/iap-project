import os
from dotenv import load_dotenv
import streamlit as st
import sqlite3
import email
from email import policy
from email.message import EmailMessage
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
import datetime
import pytz
from io import StringIO
import tempfile
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import AzureOpenAI
import json
import re
import base64  # 추가된 import

# 환경 변수 로드
load_dotenv()

# 데이터 캐시 비우기
st.cache_data.clear()
st.cache_resource.clear()

# 환경 변수에서 설정 값 직접 가져오기
STORAGE_CONN_STR = os.getenv("STORAGE_CONN_STR")
STORAGE_ACCOUNT_NAME = os.getenv("STORAGE_ACCOUNT_NAME")
EML_CONTAINER_NAME = os.getenv("EML_CONTAINER_NAME")
WORD_CONTAINER_NAME = os.getenv("WORD_CONTAINER_NAME")
DB_BASE_PATH = os.getenv("DB_BASE_PATH")

# 데이터베이스 파일 경로 구성
if DB_BASE_PATH:
    # DB_BASE_PATH가 있으면 해당 경로에 eml_reports.db 파일 생성
    EML_DB_PATH = os.path.join(DB_BASE_PATH, "eml_reports.db")
    # 디렉토리가 없으면 생성
    os.makedirs(DB_BASE_PATH, exist_ok=True)
else:
    # DB_BASE_PATH가 없으면 현재 디렉토리에 생성
    EML_DB_PATH = "eml_reports.db"

# OpenAI 설정
openai_endpoint = os.getenv("OPENAI_ENDPOINT")
openai_api_key = os.getenv("OPENAI_KEY")
openai_model = os.getenv("OPENAI_MODEL")
chat_model = os.getenv("CHAT_MODEL3")

# 환경 변수 유효성 검사
required_env_vars = {
    "STORAGE_CONN_STR": STORAGE_CONN_STR,
    "STORAGE_ACCOUNT_NAME": STORAGE_ACCOUNT_NAME,
    "EML_CONTAINER_NAME": EML_CONTAINER_NAME,
    "DB_BASE_PATH": DB_BASE_PATH
}

# 한국 시간대 설정
korea_tz = pytz.timezone('Asia/Seoul')
current_time = datetime.datetime.now(korea_tz)

# 누락된 환경 변수 확인
missing_vars = [var for var, value in required_env_vars.items() if not value]

def validate_connection_string(connection_string):
    """연결 문자열 유효성 검사"""
    if not connection_string:
        return False, "연결 문자열이 비어있습니다."
    
    # 기본적인 Azure Storage 연결 문자열 형식 확인
    required_parts = ['AccountName=', 'AccountKey=']
    missing_parts = [part for part in required_parts if part not in connection_string]
    
    if missing_parts:
        return False, f"연결 문자열에 필수 요소가 누락되었습니다: {', '.join(missing_parts)}"
    
    return True, "유효한 연결 문자열입니다."

def test_azure_connection():
    """Azure Storage 연결 테스트"""
    if not STORAGE_CONN_STR:
        return False, "연결 문자열이 설정되지 않았습니다."
    
    try:
        # 연결 문자열 유효성 검사
        is_valid, message = validate_connection_string(STORAGE_CONN_STR)
        if not is_valid:
            return False, message
        
        # 실제 연결 테스트
        blob_service_client = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
        
        # 컨테이너 존재 확인 (없으면 생성)
        try:
            container_client = blob_service_client.get_container_client(EML_CONTAINER_NAME)
            container_client.get_container_properties()
        except Exception as e:
            if "ContainerNotFound" in str(e):
                # 컨테이너가 없으면 생성
                container_client.create_container()
                return True, "연결 성공 및 컨테이너 생성 완료"
            else:
                return False, f"컨테이너 접근 오류: {str(e)}"
        
        return True, "연결 성공"
    except Exception as e:
        return False, f"연결 테스트 실패: {str(e)}"

def init_database():
    """데이터베이스 초기화 및 테이블 생성"""
    try:
        conn = sqlite3.connect(EML_DB_PATH)
        cursor = conn.cursor()
        
        # 테이블 생성 SQL
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS eml_reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                original_filename TEXT NOT NULL,
                blob_name TEXT NOT NULL,
                subject TEXT,
                sender TEXT,
                recipient TEXT,
                date_sent TEXT,
                body_text TEXT,
                body_html TEXT,
                attachments TEXT,
                file_size INTEGER,
                upload_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        conn.commit()
        conn.close()
        return True, f"데이터베이스 초기화 성공 (경로: {EML_DB_PATH})"
    except Exception as e:
        return False, f"데이터베이스 초기화 실패: {str(e)}"

def insert_eml_data(parsed_data, original_filename, blob_name, file_size):
    """EML 데이터를 데이터베이스에 삽입"""
    try:
        conn = sqlite3.connect(EML_DB_PATH)
        cursor = conn.cursor()
        
        # 첨부파일 리스트를 문자열로 변환
        attachments_str = ', '.join(parsed_data['attachments']) if parsed_data['attachments'] else ''
        
        # 데이터 삽입
        cursor.execute('''
            INSERT INTO eml_reports (
                original_filename, blob_name, subject, sender, recipient, 
                date_sent, body_text, body_html, attachments, file_size
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            original_filename,
            blob_name,
            parsed_data['subject'],
            parsed_data['from'],
            parsed_data['to'],
            parsed_data['date'],
            parsed_data['body_text'],
            parsed_data['body_html'],
            attachments_str,
            file_size
        ))
        
        record_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        return True, record_id, None
    except Exception as e:
        return False, None, f"데이터베이스 삽입 실패: {str(e)}"

def get_eml_record(record_id):
    """특정 ID의 EML 레코드 조회"""
    try:
        conn = sqlite3.connect(EML_DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT * FROM eml_reports WHERE id = ?
        ''', (record_id,))
        
        record = cursor.fetchone()
        conn.close()
        
        if record:
            # 컬럼명과 함께 딕셔너리로 반환
            columns = [desc[0] for desc in cursor.description]
            return dict(zip(columns, record))
        else:
            return None
    except Exception as e:
        st.error(f"데이터베이스 조회 실패: {str(e)}")
        return None

def get_eml_records():
    """EML 레코드 조회"""
    try:
        conn = sqlite3.connect(EML_DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, original_filename, subject, body_text
            FROM eml_reports 
            WHERE body_text IS NOT NULL AND body_text != ''
            ORDER BY upload_time DESC
        ''')
        records = cursor.fetchall()
        conn.close()
        return records
    except Exception as e:
        st.error(f"DB 오류: {e}")
        return []

def display_db_record(record):
    """데이터베이스 레코드를 화면에 표시"""
    st.subheader("💾 데이터베이스 등록 정보")
    
    # 기본 정보를 두 컬럼으로 표시
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("**📋 기본 정보**")
        st.write(f"• **ID**: {record['id']}")
        st.write(f"• **원본 파일명**: {record['original_filename']}")
        st.write(f"• **저장된 파일명**: {record['blob_name']}")
        st.write(f"• **파일 크기**: {record['file_size']:,} bytes")
    
    with col2:
        st.write("**📧 이메일 정보**")
        st.write(f"• **제목**: {record['subject']}")
        st.write(f"• **발신자**: {record['sender']}")
        st.write(f"• **수신자**: {record['recipient']}")
        st.write(f"• **발송일시**: {record['date_sent']}")
    
    # 본문 내용
    if record['body_text']:
        with st.expander("📄 본문 내용 보기"):
            st.text_area("텍스트 본문", record['body_text'], height=200, disabled=True)
    
    # 첨부파일 정보
    if record['attachments']:
        st.write("**📎 첨부파일**")
        attachments = record['attachments'].split(', ')
        for attachment in attachments:
            st.write(f"  • {attachment}")
    
    # 등록 시간
    st.write(f"**🕐 등록 시간**: {record['upload_time']}")

# parse_eml_file 함수의 HTML 처리 부분 수정 (기존 라인 270-290 부근)
def parse_eml_file(eml_content):
    """EML 파일 내용을 파싱하여 구조화된 데이터로 반환 (HTML 줄바꿈 보존 개선)"""
    try:
        # EML 파일 파싱
        msg = email.message_from_string(eml_content, policy=policy.default)
        
        # 기본 헤더 정보 추출
        parsed_data = {
            'from': msg.get('From', ''),
            'to': msg.get('To', ''),
            'cc': msg.get('Cc', ''),
            'subject': msg.get('Subject', ''),
            'date': msg.get('Date', ''),
            'message_id': msg.get('Message-ID', ''),
            'body_text': '',
            'body_html': '',
            'attachments': []
        }
        
        # 본문 내용 추출
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition", ""))
                
                # 첨부파일이 아닌 경우
                if "attachment" not in content_disposition:
                    if content_type == "text/plain":
                        try:
                            parsed_data['body_text'] = part.get_content()
                        except:
                            payload = part.get_payload(decode=True)
                            if payload:
                                parsed_data['body_text'] = payload.decode('utf-8', errors='ignore')
                    elif content_type == "text/html":
                        try:
                            parsed_data['body_html'] = part.get_content()
                        except:
                            payload = part.get_payload(decode=True)
                            if payload:
                                parsed_data['body_html'] = payload.decode('utf-8', errors='ignore')
                else:
                    # 첨부파일 정보
                    filename = part.get_filename()
                    if filename:
                        parsed_data['attachments'].append(filename)
        else:
            # 단일 파트 메시지 (Base64 UTF-8 처리 개선)
            content_type = msg.get_content_type()
            encoding = msg.get('Content-Transfer-Encoding', '').lower()
            
            if content_type == "text/plain":
                try:
                    parsed_data['body_text'] = msg.get_content()
                except Exception as e:
                    print(f"get_content() 실패, 수동 디코딩 시도: {e}")
                    try:
                        if encoding == 'base64':
                            # Base64 페이로드 추출 및 정제
                            payload = msg.get_payload()
                            if payload:
                                # 줄바꿈과 공백 제거
                                clean_payload = payload.replace('\n', '').replace('\r', '').replace(' ', '')
                                # Base64 디코딩
                                decoded_bytes = base64.b64decode(clean_payload)
                                # UTF-8로 디코딩
                                parsed_data['body_text'] = decoded_bytes.decode('utf-8', errors='ignore')
                        else:
                            payload = msg.get_payload(decode=True)
                            if payload:
                                parsed_data['body_text'] = payload.decode('utf-8', errors='ignore')
                    except Exception as decode_error:
                        print(f"텍스트 디코딩 오류: {decode_error}")
                        parsed_data['body_text'] = str(msg.get_payload())
                        
            elif content_type == "text/html":
                try:
                    parsed_data['body_html'] = msg.get_content()
                except Exception as e:
                    print(f"HTML get_content() 실패, 수동 디코딩 시도: {e}")
                    try:
                        if encoding == 'base64':
                            # Base64 페이로드 추출 및 정제
                            payload = msg.get_payload()
                            if payload:
                                # 줄바꿈과 공백 제거 (더 강력한 정제)
                                clean_payload = payload.replace('\n', '').replace('\r', '').replace(' ', '')
                                # Base64 디코딩
                                decoded_bytes = base64.b64decode(clean_payload)
                                # UTF-8로 디코딩
                                parsed_data['body_html'] = decoded_bytes.decode('utf-8', errors='ignore')
                        else:
                            payload = msg.get_payload(decode=True)
                            if payload:
                                parsed_data['body_html'] = payload.decode('utf-8', errors='ignore')
                    except Exception as decode_error:
                        print(f"HTML 디코딩 오류: {decode_error}")
                        parsed_data['body_html'] = str(msg.get_payload())
        
        # 🎯 핵심 수정: HTML에서 텍스트 추출 시 줄바꿈 보존 (body_text가 비어있을 때)
        if not parsed_data['body_text'] and parsed_data['body_html']:
            try:
                import re
                html_text = parsed_data['body_html']
                
                # 🔧 개선된 HTML 처리 - 줄바꿈 보존
                # 1. HTML 구조적 줄바꿈을 실제 줄바꿈으로 변환
                html_text = re.sub(r'</p>\s*<p[^>]*>', '\n', html_text)  # </p><p> → \n
                html_text = re.sub(r'<br\s*/?>', '\n', html_text)        # <br> → \n
                html_text = re.sub(r'</div>\s*<div[^>]*>', '\n', html_text)  # </div><div> → \n
                html_text = re.sub(r'</span>\s*<span[^>]*>', '\n', html_text)  # </span><span> → \n (시간 구분자용)
                
                # 2. 그 다음에 HTML 태그 제거
                html_text = re.sub(r'<[^>]+>', '', html_text)
                
                # 3. HTML 엔티티 처리
                html_text = html_text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                
                # 4. 연속 공백을 단일 공백으로 변경하되 줄바꿈은 보존
                html_text = re.sub(r'[ \t]+', ' ', html_text)  # 탭과 공백만 정리
                html_text = re.sub(r'\n\s*\n', '\n', html_text)  # 빈 줄 정리
                html_text = html_text.strip()
                
                parsed_data['body_text'] = html_text
                print(f"HTML에서 텍스트 추출 완료 (줄바꿈 보존): {len(html_text)}자")
                
                # 디버깅: 14:32, 14:48 관련 줄 수 확인
                debug_lines = [line for line in html_text.split('\n') if '14:32' in line or '14:48' in line]
                if debug_lines:
                    print(f"🔍 14:32/14:48 관련 줄 수: {len(debug_lines)}")
                    for i, line in enumerate(debug_lines):
                        print(f"  줄 {i+1}: {line[:100]}...")
                        
            except Exception as html_error:
                print(f"HTML 텍스트 추출 오류: {html_error}")
        
        # 추출된 텍스트 확인용 로그 (디버깅용)
        if parsed_data['body_text']:
            print(f"추출된 텍스트 길이: {len(parsed_data['body_text'])}자")
            print(f"텍스트 시작: {parsed_data['body_text'][:200]}...")
        
        return parsed_data, None
    except Exception as e:
        print(f"EML 파싱 전체 오류: {str(e)}")
        return None, str(e)

def upload_to_azure_eml_blob(file_content, filename):
    """Azure Blob Storage에 파일 업로드"""
    try:
        if not STORAGE_CONN_STR:
            return False, None, "Azure Storage 연결 문자열이 설정되지 않았습니다."
        
        # 연결 문자열 유효성 검사
        is_valid, message = validate_connection_string(STORAGE_CONN_STR)
        if not is_valid:
            return False, None, f"연결 문자열 오류: {message}"
        
        # Blob Service Client 생성
        blob_service_client = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
        
        # 타임스탬프를 포함한 고유한 파일명 생성
        timestamp = datetime.datetime.now(korea_tz).strftime("%Y%m%d_%H%M%S")
        blob_name = f"{timestamp}_{filename}"
        
        # 컨테이너 존재 확인 및 생성
        try:
            container_client = blob_service_client.get_container_client(EML_CONTAINER_NAME)
            container_client.get_container_properties()
        except Exception as e:
            if "ContainerNotFound" in str(e):
                container_client.create_container()
        
        # Blob 업로드
        blob_client = blob_service_client.get_blob_client(
            container=EML_CONTAINER_NAME, 
            blob=blob_name
        )
        
        # 파일 내용이 bytes가 아닌 경우 변환
        if isinstance(file_content, str):
            file_content = file_content.encode('utf-8')
        
        blob_client.upload_blob(file_content, overwrite=True)
        
        return True, blob_name, None
    except Exception as e:
        error_message = str(e)
        if "Connection string is either blank or malformed" in error_message:
            error_message = "Azure Storage 연결 문자열이 올바르지 않습니다. 설정을 확인해주세요."
        return False, None, error_message

def display_eml_content(parsed_data):
    """파싱된 EML 데이터를 시각적으로 표시"""
    st.subheader("📧 이메일 내용 미리보기")
    
    # 제목만 표시
    st.write("**제목:**")
    st.text(parsed_data['subject'])
    
    # 본문 내용 표시
    st.write("**본문:**")
    if parsed_data['body_text']:
        st.text_area(
            "텍스트 본문", 
            parsed_data['body_text'], 
            height=300,
            disabled=True
        )
    
    if parsed_data['body_html']:
        with st.expander("HTML 본문 보기"):
            st.code(parsed_data['body_html'], language='html')
    
    # 첨부파일 정보 표시
    if parsed_data['attachments']:
        st.write("**첨부파일:**")
        for attachment in parsed_data['attachments']:
            st.write(f"📎 {attachment}")

#현재 연도 지정을 위해 추가
current_year = datetime.datetime.now(korea_tz).year

def extract_precise_data(body_text: str) -> dict:
    """EML에서 정확한 정보만 추출 (줄별 처리로 장애 조치 결과 추출 강화)"""
    try:
        client = AzureOpenAI(
            azure_endpoint=openai_endpoint,
            api_key=openai_api_key,
            api_version="2024-02-15-preview"
        )
        
        # 🎯 이장시간 추출 개선 - AI 요청 전에 정규식으로 미리 추출
        duration_extracted = extract_duration_patterns(body_text)
        
        prompt = f"""
다음 텍스트에서 정확히 확인할 수 있는 정보만 추출하세요. 불확실한 정보는 "정보없음"으로 표시하세요.
중요: 연도가 명시되지 않은 모든 날짜는 {current_year}년으로 해석하세요.

특별히 주의할 점:
1. "장애 조치 결과" 부분은 시간순으로 나열된 모든 항목을 빠짐없이 추출해야 합니다.
2. 동일한 시간(예: 9:59, 12:32)에 여러 항목이 있으면 모두 포함해야 합니다.
3. HH:MM 형식의 시간 뒤에 나오는 모든 조치 내용을 완전히 추출해야 합니다.

특히 다음 정보들을 정확히 추출하세요:
1. "대상서비스"라는 키워드 다음에 나오는 시스템명을 정확히 추출
2. "상황반장"이라는 키워드 다음에 나오는 담당자 정보 전체를 추출
3. "복구반장"이라는 키워드 다음에 나오는 담당자 정보에서:
   - 소속부서 (예 : "KTDS ICT사업본부 ICIS Tr 추진담당 유선오더통합팀")
   - 담당자명과 직급 (예: "여재윤 책임")
4. "장애현상"이라는 키워드 다음에 나오는 장애 현상 설명과 요약 정보
5. "장애시간"이라는 키워드 다음에 나오는 괄호 안의 시간 정보 (다양한 패턴 지원)
6. "장애 조치 결과"라는 키워드 다음에 나오는 장애 조치 결과 HH:MM : 내용 패턴의 반복 정보 (반드시 누락없이 전체 추출할것)
7. "장애원인"이라는 키워드 다음에 나오는 장애 원인 분석 정보

🎯 이장시간 추출 개선사항:
- 미리 추출된 이장시간: {duration_extracted}
- 이 값이 유효하면 우선 사용하고, 없으면 텍스트에서 재추출

예시:
- "ㅇ 대상서비스 : KOS-오더(KOS-Internet)" → 시스템명: "KOS-오더(KOS-Internet)"
- "상황반장 kt ds AX사업개발본부 BA컨설팅담당 ICT컨설팅팀 윤영의 책임" → 상황반장: "ktds AX사업개발본부 BA컨설팅담당 ICT컨설팅팀 윤영의 책임"
- "복구반장 ktds ICT사업본부 ICIS Tr 추진담당 유선오더통합팀 여재윤 책임" → 
  복구반장_소속: "ktds ICT사업본부 ICIS Tr 추진담당 유선오더통합팀"
  복구반장: "여재윤 책임"
  본부이하_소속 : "ICT사업본부 ICIS Tr추진담당 유선오더통합팀"
- "장애현상 : KOS-오더 서비스 장애" → 장애현상: "KOS-오더 서비스 장애"
- "장애시간 : 04/28 14:16 ~ 15:58 (102분)" → 이장시간: "(102분)" 
- "장애현상 : KOS-오더 서비스 장애로 인한 주문 접수 불가 현상 발생" → 장애_제목: "KOS-오더 서비스 장애로 인한 주문 접수 불가"
- 장애 조치 결과에서 다음과 같은 패턴들을 모두 찾아서 추출:
  "09:59 최초 발견", "10:20 영역도 세종시 여민전 서비스 앱 접속 및 충전 불가", "11:42 충전 중지없이 현상태 유지", "12:00 복구예상시간 13:00로 협의" 등
  → "장애_조치_경과_리스트": [
    {{"시간": "09:59", "내용": "최초 발견", "비고": ""}},
    {{"시간": "10:20", "내용": "영역도 세종시 여민전 서비스 앱 접속 및 충전 불가", "비고": ""}},
    {{"시간": "11:42", "내용": "충전 중지없이 현상태 유지하기로함(융합서비스플랫폼팀)", "비고": ""}},
    {{"시간": "12:00", "내용": "복구예상시간 13:00로 세종시에 협의됨(융합서비스플랫폼팀,결제/센싱서비스팀)", "비고": ""}}
]

{body_text}

JSON 응답:
{{
    "장애_제목": "장애현상 요약 정보(문장 뒤쪽 현상, 발생 등을 제외한 요약 정보)",
    "시스템명": "대상서비스 뒤에 나오는 정확한 시스템명",
    "장애_등급": "명시된 등급만",
    "발생_시간": "정확한 날짜/시간만 ({current_year}년 MM월 DD일 HH:MM 형식)",
    "인지_시간": "정확한 인지 시간만",
    "복구_시간": "정확한 복구 시간만 ({current_year}년 MM월 DD일 HH:MM 형식)",
    "이장_시간": "정확한 이장 시간만 - 미리 추출된 값 우선 사용: {duration_extracted}",
    "장애_현상": "명확한 현상 설명만",
    "파급_영향": "명확한 영향 설명만",
    "근본_원인": "명확한 원인 분석만",
    "상황반장": "상황반장 전체 정보",
    "복구반장_소속": "복구반장의 소속부서만 (복구반장의 이름과 직급 제외)",
    "복구반장": "복구반장의 이름과 직급만",
    "본부이하_소속": "회사(KTDS)를 제외한 복구반장 소속 정보(ㅇㅇㅇ본부 ㅇㅇㅇ담당 ㅇㅇㅇ팀)",
    "장애_조치_경과_리스트": [
        {{"시간": "HH:MM", "내용": "작업 내용/현상", "비고": ""}},
        {{"시간": "HH:MM", "내용": "작업 내용/현상", "비고": ""}}
    ]
}}

불확실하면 "정보없음"으로 응답하세요.
"""
        
        response = client.chat.completions.create(
            model=chat_model,
            messages=[
                {"role": "system", "content": "정확한 정보만 추출하는 전문가. 특히 '장애 조치 결과' 부분에서 시간순으로 나열된 모든 항목을 빠뜨리지 않고 추출하며, 동일 시간대의 여러 항목도 모두 포함함. '대상서비스', '상황반장', '복구반장', '장애현상' 등 키워드 뒤의 정보를 정확히 찾아서 추출함."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.0
        )
        
        result = response.choices[0].message.content.strip()
        if "```json" in result:
            start = result.find("```json") + 7
            end = result.rfind("```")
            result = result[start:end].strip()
        
        extracted_data = json.loads(result)
        
        # 🎯 이장시간 후처리 - AI가 추출하지 못했거나 잘못 추출한 경우 정규식 결과로 보완
        if duration_extracted and (not extracted_data.get("이장_시간") or extracted_data.get("이장_시간") == "정보없음"):
            extracted_data["이장_시간"] = duration_extracted
            if 'extract_logs' not in st.session_state:
                st.session_state.extract_logs = []
            st.session_state.extract_logs.append(f"🎯 정규식으로 이장시간 보완: {duration_extracted}")
        
        # 정규식을 이용한 강화된 장애 조치 결과 추출 (백업) - 🎯 줄별 처리로 개선 
        if 'extract_logs' not in st.session_state:
            st.session_state.extract_logs = []
        
        st.session_state.extract_logs = []  # 로그 초기화
        
        # 장애 조치 결과 정규식 추출 (강화된 버전 - 줄별 처리)
        action_pattern = r'장애\s*조치\s*결과\s*[:：]\s*(.+?)(?:\n\n|\n[가-힣]+\s*[:：]|장애원인|$)'
        action_match = re.search(action_pattern, body_text, re.DOTALL)
        
        if action_match:
            action_text = action_match.group(1).strip()
            st.session_state.extract_logs.append(f"🎯 조치 결과 텍스트 추출: {len(action_text)}자")
            
            # 🔧 개선된 줄별 처리 방식
            action_list_regex = []
            
            # 줄바꿈으로 분할하여 각 줄을 개별 처리
            lines = action_text.split('\n')
            cleaned_lines = [line.strip() for line in lines if line.strip()]
            
            st.session_state.extract_logs.append(f"🔍 전체 줄 수: {len(cleaned_lines)}")
            
            # 각 줄에서 시간:내용 패턴 찾기
            for line_idx, line in enumerate(cleaned_lines):
                # 시간 패턴 매칭 (줄 시작부터)
                time_match = re.match(r'(\d{2}:\d{2})\s*(.+)', line)
                if time_match:
                    time_str = time_match.group(1)
                    content = time_match.group(2).strip()
                    
                    # 내용 정리 (불필요한 문자 제거)
                    content = content.replace(':', '').replace('：', '').strip()
                    
                    if content and len(content) > 2:
                        action_list_regex.append({
                            "시간": time_str,
                            "내용": content,
                            "비고": ""
                        })
                        st.session_state.extract_logs.append(f"✅ 줄 {line_idx+1}: {time_str} - {content[:50]}...")
                    else:
                        st.session_state.extract_logs.append(f"⚠️ 줄 {line_idx+1}: {time_str} - 내용 부족 ({content})")
            
            # 14:32, 14:48 관련 디버깅 정보
            debug_lines = [line for line in cleaned_lines if '14:32' in line or '14:48' in line]
            if debug_lines:
                st.session_state.extract_logs.append(f"🔍 14:32/14:48 관련 줄 수: {len(debug_lines)}")
                for i, line in enumerate(debug_lines):
                    st.session_state.extract_logs.append(f"  디버그 줄 {i+1}: {line}")
            
            # AI 추출 결과와 비교
            ai_action_count = 0
            if extracted_data.get("장애_조치_경과_리스트") and extracted_data["장애_조치_경과_리스트"] != "정보없음":
                ai_action_count = len(extracted_data["장애_조치_경과_리스트"])
            
            regex_action_count = len(action_list_regex)
            
            st.session_state.extract_logs.append(f"🔍 AI 추출: {ai_action_count}개, 정규식 추출 (줄별): {regex_action_count}개")
            
            # 정규식 결과가 더 많거나 같으면 정규식 결과 사용
            if regex_action_count >= ai_action_count:
                extracted_data["장애_조치_경과_리스트"] = action_list_regex
                st.session_state.extract_logs.append(f"✅ 정규식 결과 사용 (줄별 처리): {regex_action_count}개 항목")
        
        # 추가적으로 정규식으로 직접 추출 (기존 백업 코드들)
        # 대상서비스 추출
        service_pattern = r'대상서비스\s*[:：]\s*(.+?)(?:\n|$)'
        service_match = re.search(service_pattern, body_text)
        if service_match:
            service_name = service_match.group(1).strip()
            if service_name and extracted_data.get("시스템명") == "정보없음":
                extracted_data["시스템명"] = service_name
                st.session_state.extract_logs.append(f"🎯 정규식으로 시스템명 추출: {service_name}")
        
        # 상황반장 추출
        situation_pattern = r'상황반장\s*[:：]?\s*(.+?)(?:\n|복구반장|$)'
        situation_match = re.search(situation_pattern, body_text)
        if situation_match:
            situation_leader = situation_match.group(1).strip()
            if situation_leader and extracted_data.get("상황반장") == "정보없음":
                extracted_data["상황반장"] = situation_leader
                st.session_state.extract_logs.append(f"🎯 정규식으로 상황반장 추출: {situation_leader}")
        
        # 복구반장 추출
        recovery_pattern = r'복구반장\s*[:：]?\s*(.+?)(?:\n|$)'
        recovery_match = re.search(recovery_pattern, body_text)
        if recovery_match:
            recovery_full = recovery_match.group(1).strip()
            if recovery_full:
                # 복구반장 정보를 소속과 담당자로 분리
                # 일반적으로 마지막 두 단어가 "이름 직급" 형태
                words = recovery_full.split()
                if len(words) >= 2:
                    # 마지막 두 단어를 담당자로, 나머지를 소속으로
                    recovery_person = " ".join(words[-2:])
                    recovery_dept = " ".join(words[:-2])
                    
                    if extracted_data.get("복구반장_소속") == "정보없음":
                        extracted_data["복구반장_소속"] = recovery_dept
                    if extracted_data.get("복구반장") == "정보없음":
                        extracted_data["복구반장"] = recovery_person
                    
                    st.session_state.extract_logs.append(f"🎯 정규식으로 복구반장 분리: 소속={recovery_dept}, 담당자={recovery_person}")
        
        # 장애현상 추출
        incident_pattern = r'장애현상\s*[:：]\s*(.+?)(?:\n|$)'
        incident_match = re.search(incident_pattern, body_text)
        if incident_match:
            incident_symptom = incident_match.group(1).strip()
            if incident_symptom and extracted_data.get("장애현상") == "정보없음":
                extracted_data["장애현상"] = incident_symptom
                st.session_state.extract_logs.append(f"🎯 정규식으로 장애현상 추출: {incident_symptom}")

        return extracted_data
        
    except Exception as e:
        if 'extract_logs' not in st.session_state:
            st.session_state.extract_logs = []
        st.session_state.extract_logs.append(f"❌ AI 추출 오류: {e}")
        return {"error": "추출 실패"}
    
def extract_duration_patterns(body_text: str) -> str:
    """다양한 이장시간 패턴을 추출하는 함수"""
    try:
        # 🎯 개선된 이장시간 추출 패턴들
        patterns = [
            # 1. 괄호 안에 숫자와 분이 있는 패턴 (가장 일반적)
            r'\((\d+분)\)',                           # (6분)
            r'\(이장시간\s*:?\s*(\d+분)\)',           # (이장시간 6분), (이장시간: 6분)
            r'\(.*?(\d+분).*?\)',                     # (기타 텍스트 6분 기타)
            
            # 2. 시간 범위에서 분 계산 (HH:MM ~ HH:MM 형태)
            r'(\d{2}:\d{2})\s*~\s*(\d{2}:\d{2})',    # 09:51 ~ 09:57
            
            # 3. 장애시간 키워드 뒤의 괄호 패턴
            r'장애시간.*?\(.*?(\d+분).*?\)',          # 장애시간 : ... (102분)
            r'이장시간.*?\(.*?(\d+분).*?\)',          # 이장시간 : ... (6분)
            
            # 4. 기타 시간 표현
            r'(\d+)분\s*간',                         # 6분간
            r'이\s*(\d+)분',                         # 이 6분
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, body_text, re.IGNORECASE)
            if matches:
                if len(matches[0]) == 2 and ':' in matches[0][0]:
                    # 시간 범위 패턴인 경우 (HH:MM ~ HH:MM)
                    start_time, end_time = matches[0]
                    try:
                        start_h, start_m = map(int, start_time.split(':'))
                        end_h, end_m = map(int, end_time.split(':'))
                        
                        start_total_minutes = start_h * 60 + start_m
                        end_total_minutes = end_h * 60 + end_m
                        
                        # 시간이 다음날로 넘어간 경우 처리
                        if end_total_minutes < start_total_minutes:
                            end_total_minutes += 24 * 60
                        
                        duration_minutes = end_total_minutes - start_total_minutes
                        return f"({duration_minutes}분)"
                    except:
                        continue
                else:
                    # 다른 패턴들의 경우
                    duration = matches[0] if isinstance(matches[0], str) else matches[0][0]
                    if '분' in duration:
                        return f"({duration})"
                    else:
                        return f"({duration}분)"
        
        return ""  # 패턴을 찾지 못한 경우
        
    except Exception as e:
        print(f"이장시간 추출 오류: {e}")
        return ""

def find_action_progress_table(doc):
    """장애 조치 경과 표 찾기"""
    try:
        # 로그 저장을 위한 session_state 초기화
        if 'template_logs' not in st.session_state:
            st.session_state.template_logs = []
        
        for table_idx, table in enumerate(doc.tables):
            # 첫 번째 행에서 "일시", "작업 내용", "비고" 같은 헤더 찾기
            if len(table.rows) > 0:
                header_row = table.rows[0]
                header_text = " ".join([cell.text.strip() for cell in header_row.cells])
                
                # 조치 경과 관련 헤더 키워드들
                progress_keywords = ["일시", "시간", "작업", "내용", "현상", "비고", "경과"]
                matched_count = sum(1 for keyword in progress_keywords if keyword in header_text)
                
                if matched_count >= 2:  # 최소 2개 키워드 매칭되면 조치 경과 표로 판단
                    st.session_state.template_logs.append(f"🎯 조치 경과 표 발견: 테이블 {table_idx} (헤더: {header_text})")
                    return table_idx, table
                    
        # 키워드로 찾지 못했다면 "조치" 키워드가 포함된 표 찾기
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    if "조치" in cell.text and ("경과" in cell.text or "결과" in cell.text):
                        st.session_state.template_logs.append(f"🎯 조치 경과 표 발견 (키워드 매칭): 테이블 {table_idx}")
                        return table_idx, table
                        
        return None, None
        
    except Exception as e:
        if 'template_logs' not in st.session_state:
            st.session_state.template_logs = []
        st.session_state.template_logs.append(f"⚠️ 조치 경과 표 찾기 오류: {e}")
        return None, None

def fill_action_progress_table(table, action_list):
    """조치 경과 표에 데이터 입력"""
    if not action_list or not table:
        return 0
    
    # 로그 저장을 위한 session_state 초기화
    if 'template_logs' not in st.session_state:
        st.session_state.template_logs = []
    
    filled_count = 0
    
    try:
        # 헤더 행은 건드리지 않고, 두 번째 행부터 데이터 입력
        start_row = 1
        
        # 필요한 행 수 계산
        needed_rows = len(action_list)
        current_rows = len(table.rows)
        
        # 행이 부족하면 추가
        if current_rows < needed_rows + start_row:
            rows_to_add = needed_rows + start_row - current_rows
            st.session_state.template_logs.append(f"🔍 조치 경과 표에 {rows_to_add}개 행 추가")
            
            for _ in range(rows_to_add):
                # 새 행 추가 (첫 번째 행의 셀 수만큼)
                new_row = table.add_row()
                filled_count += 1
        
        # 데이터 입력
        for idx, action_item in enumerate(action_list):
            row_idx = start_row + idx
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                
                # 각 셀에 데이터 입력 (시간, 내용, 비고 순서)
                if len(row.cells) >= 1:
                    row.cells[0].text = action_item.get("시간", "")
                if len(row.cells) >= 2:
                    row.cells[1].text = action_item.get("내용", "")
                if len(row.cells) >= 3:
                    row.cells[2].text = action_item.get("비고", "")
                
                filled_count += 1
                st.session_state.template_logs.append(f"✅ {action_item.get('시간')} - {action_item.get('내용')}")
        
        return filled_count
        
    except Exception as e:
        st.session_state.template_logs.append(f"❌ 조치 경과 표 입력 오류: {e}")
        return 0

def replace_placeholder_text(doc, placeholder, replacement):
    """문서 전체에서 플레이스홀더를 찾아서 교체"""
    replaced_count = 0
    
    # 문단에서 교체
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            replaced_count += 1
    
    # 표에서 교체
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)
                    replaced_count += 1
    
    return replaced_count

def fill_template_safely(template_path: str, data: dict, record_id: int) -> str:
    """안전한 템플릿 채우기 (개선된 키워드 기반 방식)"""
    try:
        # 로그 저장을 위한 session_state 초기화
        if 'template_logs' not in st.session_state:
            st.session_state.template_logs = []
        
        st.session_state.template_logs = []  # 로그 초기화
        
        doc = Document(template_path)
        filled_count = 0
        
        # 1. 시스템명 플레이스홀더 교체 (#서비스명 → KOS-오더(KOS-Internet))
        system_name = data.get("시스템명", "")
        if system_name and system_name != "정보없음":
            replaced = replace_placeholder_text(doc, "(#서비스명)", system_name)
            if replaced > 0:
                filled_count += replaced
                st.session_state.template_logs.append(f"🎯 시스템명 교체 완료: (#서비스명) → {system_name} ({replaced}곳)")
            else:
                st.session_state.template_logs.append("⚠️ (#서비스명) 플레이스홀더를 찾을 수 없습니다")
        
        # 2. 기타 플레이스홀더들 교체 (상황반장 포함)
        placeholder_mappings = {
            "(#장애등급)": data.get("장애_등급", ""),
            "(#발생시간)": data.get("발생_시간", ""),
            "(#인지시간)": data.get("인지_시간", ""),
            "(#복구시간)": data.get("복구_시간", ""),
            "(#이장시간)": data.get("이장_시간", ""),
            "(#소속본부)": data.get("소속_본부", ""),
            "(#소속팀)": data.get("소속_팀", ""),
            "(#상황반장)": data.get("상황반장", ""),
            "(#복구반장 소속)": data.get("복구반장_소속", ""),
            "(#보고자 소속)": data.get("본부이하_소속", ""),
            "(#복구반장)": data.get("복구반장", ""),
            "(#장애현상)": data.get("장애_현상", ""),
            "(#장애파급영향)": data.get("파급_영향", ""),
            "(#장애근본원인)": data.get("근본_원인", ""),
        }
        
        for placeholder, value in placeholder_mappings.items():
            if value and value != "정보없음":
                replaced = replace_placeholder_text(doc, placeholder, value)
                if replaced > 0:
                    filled_count += replaced
                    st.session_state.template_logs.append(f"✅ {placeholder} → {value} ({replaced}곳)")
        
        # 상황반장 별도 처리 (플레이스홀더와 키워드 매칭 모두 시도)
        situation_leader = data.get("상황반장", "")
        if situation_leader and situation_leader != "정보없음":
            # 1. 플레이스홀더 교체 시도
            replaced = replace_placeholder_text(doc, "(#상황반장)", situation_leader)
            if replaced > 0:
                filled_count += replaced
                st.session_state.template_logs.append(f"🎯 상황반장 플레이스홀더 교체: {situation_leader} ({replaced}곳)")
                    
        # 3. 운영부서에도 복구반장 소속 정보 입력 (추가 보장)
        recovery_dept = data.get("복구반장_소속", "")
        if recovery_dept and recovery_dept != "정보없음":
            # 운영부서 플레이스홀더도 같은 값으로 교체
            replaced = replace_placeholder_text(doc, "(#운영부서)", recovery_dept)
            if replaced > 0:
                filled_count += replaced
                st.session_state.template_logs.append(f"✅ (#운영부서) → {recovery_dept} ({replaced}곳)")

        # 4. 제목 입력 (첫 번째 문단에 강제 입력)
        title = data.get("장애_제목", "")
        if title and title != "정보없음":
            # 첫 번째 문단에 제목 입력 (기존 내용 덮어쓰기)
            if len(doc.paragraphs) > 0:
                doc.paragraphs[0].text = title
                filled_count += 1
                st.session_state.template_logs.append(f"🎯 제목 입력: {title}")
            else:
                # 문단이 없으면 새로 추가
                new_para = doc.add_paragraph(title)
                doc._body._element.insert(0, new_para._element)
                filled_count += 1
                st.session_state.template_logs.append(f"🎯 제목 추가: {title}")
        
        # 5. 소속 정보 입력 (두 번째 문단) - 수정된 부분
        dept_text = ""
        if data.get("본부이하_소속") and data.get("본부이하_소속") != "정보없음":
            dept_text = data.get("본부이하_소속")  # 공백 제거
        dept_text += f" ({current_time.strftime('%Y.%m.%d')})"
        
        # 두 번째 내용이 있는 문단에 입력 (서식 유지)
        content_paras = [p for p in doc.paragraphs if p.text.strip()]
        if len(content_paras) >= 2:
            # 기존 서식을 유지하면서 텍스트만 교체
            paragraph = content_paras[1]
            
            # 기존 run들을 모두 제거하고 새로운 run 추가
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # 새로운 run 추가 (기본 서식 적용)
            new_run = paragraph.add_run(dept_text)
            
            # 폰트 크기와 굵기 설정 (템플릿과 동일하게)
            new_run.font.size = Pt(10)  # 폰트 크기 10pt
            new_run.font.bold = False   # 굵기 제거
            
            filled_count += 1
            st.session_state.template_logs.append(f"🎯 소속 정보 입력: {dept_text} (서식 적용)")
        
        # 6. 장애 조치 경과 표에 데이터 입력
        action_list = data.get("장애_조치_경과_리스트", [])
        if action_list and action_list != "정보없음":
            table_idx, action_table = find_action_progress_table(doc)
            if action_table:
                action_filled = fill_action_progress_table(action_table, action_list)
                filled_count += action_filled
                st.session_state.template_logs.append(f"🎯 조치 경과 표 입력 완료: {len(action_list)}개 항목, {action_filled}개 행 처리")
            else:
                st.session_state.template_logs.append("⚠️ 조치 경과 표를 찾을 수 없습니다")
            
        # 7. 생성 정보 추가
        # doc.add_paragraph(f"\n[AI 생성: {current_time.strftime('%Y-%m-%d %H:%M:%S')} | 레코드: {record_id} | 입력: {filled_count}개]")
        
        # 8. 저장
        temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        doc.save(temp_path)
        
        st.session_state.template_logs.append(f"✅ 총 {filled_count}개 필드 입력 완료")
        return temp_path
        
    except Exception as e:
        if 'template_logs' not in st.session_state:
            st.session_state.template_logs = []
        st.session_state.template_logs.append(f"❌ 템플릿 처리 오류: {e}")
        raise Exception(f"템플릿 채우기 실패: {e}")

def upload_to_azure_word(file_path: str, filename: str):
    """24시간 유효 Azure 업로드"""
    try:
        blob_service_client = BlobServiceClient.from_connection_string(STORAGE_CONN_STR)
        timestamp = datetime.datetime.now(korea_tz).strftime("%Y%m%d_%H%M%S")
        blob_name = f"{timestamp}_{filename}"
        
        # 업로드
        blob_client = blob_service_client.get_blob_client(container=WORD_CONTAINER_NAME, blob=blob_name)
        with open(file_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
        
        # 24시간 SAS 토큰
        account_key = STORAGE_CONN_STR.split('AccountKey=')[1].split(';')[0]
        sas_token = generate_blob_sas(
            account_name=STORAGE_ACCOUNT_NAME,
            container_name=WORD_CONTAINER_NAME,
            blob_name=blob_name,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.datetime.now(korea_tz) + datetime.timedelta(hours=24)
        )
        
        url = f"https://{STORAGE_ACCOUNT_NAME}.blob.core.windows.net/{WORD_CONTAINER_NAME}/{blob_name}?{sas_token}"
        return True, url, None
        
    except Exception as e:
        return False, None, str(e)
    
##보고서(초안) 활용 가이드
def show_completion_guide_simple():
    """간단한 완료 가이드 표시"""
    
    # 주의사항
    st.warning("⚠️ **중요 안내**\n\n보고서는 첨부 복구보고(eml) 정보를 기반으로 AI가 자동 생성한 초안입니다. 초안에 오류가 있을 수 있으니 모든 항목을 사용자께서 직접 확인하시기 바랍니다.")
        
    # 가이드 다운로드
    st.subheader("📋 보고서 활용 가이드")
    st.info("Best Practice를 제공하니 첨부 파일을 참고하세요.")
    
    # 로컬 파일 경로
    guide_path = "data/docx/20250320_장애보고서_KT AICC 공공지자체_고객포털접속지연장애_v1.0.docx"
    guide_exists = os.path.exists(guide_path)
    
    col1, col2 = st.columns(2)
    with col1:
        if guide_exists:
            with open(guide_path, "rb") as file:
                st.download_button(
                    label="📥 장애원인분석보고서(B·P) 다운로드",
                    data=file.read(),
                    file_name="20250320_장애보고서_KT AICC 공공지자체_고객포털접속지연장애_v1.0.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("⚠️ 가이드 파일을 찾을 수 없습니다.")
    
    with col2:
        if st.button("🔍 미리보기"):
            # 미리보기 이미지 (로컬 파일이 있다고 가정)
            preview_image_path = "data/docx/guide-preview.png"
            if os.path.exists(preview_image_path):
                st.image(preview_image_path)
    st.divider()


# === Streamlit 앱 ===
st.title("💡 장애보고서 초안 생성기")
st.caption("복구보고(eml)을 업로드하면 장애보고서 초안을 생성 합니다")

# 데이터베이스 초기화
db_success, db_message = init_database()

# Azure 연결 상태 확인
connection_test_result, connection_message = test_azure_connection()

# 전역 상태 초기화
if 'stage' not in st.session_state:
    st.session_state.stage = 'upload'  # 'upload', 'processing', 'completed'

if 'current_record_id' not in st.session_state:
    st.session_state.current_record_id = None

if 'extracted' not in st.session_state:
    st.session_state.extracted = None

# 템플릿 파일 확인
template_path = "data/docx/iap-report-sample1(#).docx"
template_exists = os.path.exists(template_path)

# 1단계: EML 파일 업로드
if st.session_state.stage == 'upload':
    st.subheader("📁 EML 파일 업로드")
    
    uploaded_file = st.file_uploader(
        "복구보고서 EML 파일을 선택하세요",
        type=['eml'],
        help="복구 보고서가 포함된 EML 파일을 업로드해주세요"
    )
    
    if uploaded_file is not None:
        # 파일 내용 읽기
        file_content = uploaded_file.read()
        file_content_str = file_content.decode('utf-8', errors='ignore')
        
        # EML 파일 파싱
        with st.spinner('EML 파일을 분석 중입니다...'):
            parsed_data, parse_error = parse_eml_file(file_content_str)
        
        if parse_error:
            st.error(f"❌ EML 파일 파싱 중 오류가 발생했습니다: {parse_error}")
        else:
            # 파싱된 내용 표시
            display_eml_content(parsed_data)
            
            st.divider()
            
            # 업로드 확인
            if not connection_test_result:
                st.error("❌ Azure Storage 연결이 설정되지 않아 업로드할 수 없습니다.")
            elif not db_success:
                st.error("❌ 데이터베이스 연결이 설정되지 않아 업로드할 수 없습니다.")
            else:
                if st.button("✅ 업로드 및 다음 단계", type="primary"):
                    with st.spinner('Azure Storage에 업로드 중입니다...'):
                        success, blob_name, upload_error = upload_to_azure_eml_blob(
                            file_content, 
                            uploaded_file.name
                        )
                    
                    if success:
                        st.success(f"✅ 파일이 성공적으로 업로드되었습니다!")
                        
                        # 데이터베이스에 정보 저장
                        with st.spinner('데이터베이스에 정보를 저장 중입니다...'):
                            db_success_insert, record_id, db_error = insert_eml_data(
                                parsed_data, 
                                uploaded_file.name, 
                                blob_name, 
                                len(file_content)
                            )
                        
                        if db_success_insert:
                            st.success(f"✅ 데이터베이스에 성공적으로 등록되었습니다! (ID: {record_id})")
                            
                            # 세션 상태 업데이트
                            st.session_state.current_record_id = record_id
                            st.session_state.current_filename = uploaded_file.name
                            st.session_state.current_body_text = parsed_data['body_text']
                            st.session_state.stage = 'processing'
                            st.rerun()
                        else:
                            st.error(f"❌ 데이터베이스 등록 중 오류가 발생했습니다: {db_error}")
                    else:
                        st.error(f"❌ 업로드 중 오류가 발생했습니다: {upload_error}")

# 2단계: AI 분석 및 보고서 생성
elif st.session_state.stage == 'processing':
    st.subheader("🔍 AI 분석 및 보고서 생성")
    
    # 현재 파일 정보 표시
    st.info(f"📄 분석 대상: {st.session_state.current_filename} (ID: {st.session_state.current_record_id})")
    
    # OpenAI 설정 확인
    if not openai_endpoint or not openai_api_key:
        st.error("❌ OpenAI 설정이 완료되지 않았습니다.")
        st.stop()
    
    # 템플릿 파일 확인
    if not template_exists:
        st.error("❌ 템플릿 파일을 찾을 수 없습니다.")
        st.stop()
    
    # AI 분석 실행
    if st.session_state.extracted is None:
        with st.spinner("AI가 장애보고서 정보를 추출 중입니다..."):
            extracted = extract_precise_data(st.session_state.current_body_text)
            if "error" not in extracted:
                st.session_state.extracted = extracted
                st.success("✅ 정보 추출 완료!")
                
                # AI 분석 로그 표시
                if hasattr(st.session_state, 'extract_logs') and st.session_state.extract_logs:
                    with st.expander("🔍 AI 분석 로그 보기"):
                        for log in st.session_state.extract_logs:
                            st.write(log)
            else:
                st.error("❌ 정보 추출에 실패했습니다.")
                if st.button("🔄 다시 시도"):
                    st.rerun()
                st.stop()
    
    # 추출된 정보 표시
    if st.session_state.extracted:
        data = st.session_state.extracted
        valid_data = {k: v for k, v in data.items() if v and v != "정보없음"}
        
        st.subheader("📋 추출된 정보")
        with st.expander("추출된 데이터 보기"):
            st.json(valid_data)
        
        st.divider()
        
        # 보고서 생성
        st.subheader("📄 장애분석보고서 생성")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 보고서 생성", type="primary"):
                with st.spinner("보고서를 생성 중입니다..."):
                    try:
                        # 보고서 생성
                        report_path = fill_template_safely(template_path, data, st.session_state.current_record_id)
                        
                        # 템플릿 처리 로그 표시
                        if hasattr(st.session_state, 'template_logs') and st.session_state.template_logs:
                            with st.expander("🔍 템플릿 처리 로그 보기"):
                                for log in st.session_state.template_logs:
                                    st.write(log)
                        
                        # Azure 업로드
                        success, url, error = upload_to_azure_word(report_path, f"보고서_{st.session_state.current_record_id}.docx")
                        os.unlink(report_path)
                        
                        if success:
                            st.success("✅ 보고서 생성 완료!")
                            st.session_state.report_url = url
                            st.session_state.stage = 'completed'
                            st.rerun()
                        else:
                            st.error(f"❌ 보고서 업로드 실패: {error}")
                            
                    except Exception as e:
                        st.error(f"❌ 보고서 생성 실패: {e}")
        
        with col2:
            if st.button("🔄 새 파일 업로드"):
                # 상태 초기화
                st.session_state.stage = 'upload'
                st.session_state.current_record_id = None
                st.session_state.extracted = None
                st.rerun()

# 3단계: 완료
elif st.session_state.stage == 'completed':
    st.subheader("✅ 보고서 생성 완료")
    
    st.success(f"장애분석보고서가 성공적으로 생성되었습니다!")
    
    # 다운로드 링크
    if hasattr(st.session_state, 'report_url'):
        st.markdown(f"### [📥 다운로드 - AI 생성 보고서]({st.session_state.report_url})")
        st.info("💡 다운로드 링크는 24시간 동안 유효합니다.")
    
    st.divider()
    
    # 🎯 가이드 표시
    show_completion_guide_simple()

    # 추가 작업 옵션
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔄 새 보고서 생성"):
            # 상태 초기화
            st.session_state.stage = 'upload'
            st.session_state.current_record_id = None
            st.session_state.extracted = None
            if hasattr(st.session_state, 'report_url'):
                delattr(st.session_state, 'report_url')
            st.rerun()
    
    with col2:
        if st.button("📋 기존 파일로 다시 생성"):
            # 같은 파일로 다시 생성
            st.session_state.stage = 'processing'
            st.session_state.extracted = None  # 재분석
            st.rerun()