import os
from dotenv import load_dotenv
import streamlit as st
import sqlite3
from docx import Document
from docx.shared import Pt
import tempfile
from azure.storage.blob import BlobServiceClient
import datetime
from openai import AzureOpenAI

# 환경 변수 로드
load_dotenv()

# 데이터 캐시 비우기
st.cache_data.clear()
st.cache_resource.clear()


# Azure OpenAI 설정
openai_endpoint = os.getenv("OPENAI_ENDPOINT")
openai_api_key = os.getenv("OPENAI_KEY")
chat_model = os.getenv("CHAT_MODEL")
# 환경변수 설정
AZURE_STORAGE_CONNECTION_STRING = os.getenv("STORAGE_CONN_STR")

STORAGE_ACCOUNT_NAME = os.getenv("STORAGE_ACCOUNT_NAME")
WORD_CONTAINER_NAME = os.getenv("WORD_CONTAINER_NAME")
EML_DB_NAME = os.getenv("EML_DB_NAME")


def get_all_eml_records():
    """모든 EML 레코드 조회"""
    try:
        conn = sqlite3.connect(EML_DB_NAME)
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, original_filename, subject, body_text
            FROM eml_reports 
            WHERE body_text IS NOT NULL AND body_text != ''
            ORDER BY upload_time DESC
        ''')
        
        records = cursor.fetchall()
        conn.close()
        
        return records
    except Exception as e:
        st.error(f"데이터베이스 조회 실패: {str(e)}")
        return []

def read_sample_document(file_path: str) -> str:
    """샘플 문서의 내용을 읽어서 텍스트로 반환"""
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 빈 문단 제외
                full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"샘플 문서를 읽을 수 없습니다: {str(e)}"

def generate_llm_report(body_text: str, sample_format: str) -> str:
    """Azure OpenAI를 사용하여 LLM 보고서 생성"""
    
    try:
        if not all([openai_endpoint, openai_api_key, chat_model]):
            return "Azure OpenAI 설정이 완료되지 않았습니다. 환경변수를 확인해주세요."
        
        # Azure OpenAI 클라이언트 초기화
        client = AzureOpenAI(
            azure_endpoint=openai_endpoint,
            api_key=openai_api_key,
            api_version="2024-02-15-preview"
        )
        
        # 프롬프트 생성
        prompt = f"""
다음 샘플 장애보고서의 형식을 참고하여, 주어진 본문 내용을 바탕으로 전문적인 장애보고서를 작성해주세요.

샘플 장애보고서 형식:
{sample_format}

본문 내용:
{body_text}

위의 본문 내용을 분석하여 샘플 형식에 맞는 전문적인 장애보고서를 작성해주세요. 
- 장애 발생 시간, 원인, 영향도, 조치사항 등을 명확히 구분하여 작성
- 기술적인 내용은 정확하고 이해하기 쉽게 설명
- 보고서 형식은 샘플과 동일하게 유지
"""
        
        # OpenAI API 호출
        response = client.chat.completions.create(
            model=chat_model,
            messages=[
                {"role": "system", "content": "당신은 IT 장애보고서 작성 전문가입니다. 주어진 정보를 바탕으로 정확하고 전문적인 장애보고서를 작성해주세요."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.3
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        return f"LLM 보고서 생성 중 오류가 발생했습니다: {str(e)}"

def generate_word_from_body(body_text: str, filename: str) -> str:
    """본문 내용을 기반으로 Word 문서 생성"""
    doc = Document("data/docx/iap-report-sample1.docx")  # 샘플 문서 기반
    
    doc.add_paragraph("")  # 구분선
    content_paragraph = doc.add_paragraph(body_text)
    content_paragraph.style.font.size = Pt(11)

    # 임시 파일에 저장
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_path)
    return temp_path

def upload_to_azure_word_blob(file_content, filename):
    """Azure Blob Storage에 파일 업로드"""
    try:
        if not AZURE_STORAGE_CONNECTION_STRING:
            return False, None, "Azure Storage 연결 문자열이 설정되지 않았습니다."
        
        # 연결 문자열 유효성 검사
        is_valid, message = validate_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        if not is_valid:
            return False, None, f"연결 문자열 오류: {message}"
        
        # Blob Service Client 생성
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        
        # 타임스탬프를 포함한 고유한 파일명 생성
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        blob_name = f"{timestamp}_{filename}"
        
        # 컨테이너 존재 확인 및 생성
        try:
            container_client = blob_service_client.get_container_client(WORD_CONTAINER_NAME)
            container_client.get_container_properties()
        except Exception as e:
            if "ContainerNotFound" in str(e):
                container_client.create_container()
        
        # Blob 업로드
        blob_client = blob_service_client.get_blob_client(
            container=WORD_CONTAINER_NAME, 
            blob=blob_name
        )
        
        # 파일 내용이 bytes가 아닌 경우 변환
        if isinstance(file_content, str):
            file_content = file_content.encode('utf-8')
        
        blob_client.upload_blob(file_content, overwrite=True)
        
        return True, blob_name, None
    except Exception as e:
        error_message = str(e)
        if "Connection string is either blank or malformed" in error_message:
            error_message = "Azure Storage 연결 문자열이 올바르지 않습니다. 설정을 확인해주세요."
        return False, None, error_message

def validate_connection_string(connection_string):
    """연결 문자열 유효성 검사"""
    if not connection_string:
        return False, "연결 문자열이 비어있습니다."
    
    # 기본적인 Azure Storage 연결 문자열 형식 확인
    required_parts = ['AccountName=', 'AccountKey=']
    missing_parts = [part for part in required_parts if part not in connection_string]
    
    if missing_parts:
        return False, f"연결 문자열에 필수 요소가 누락되었습니다: {', '.join(missing_parts)}"
    
    return True, "유효한 연결 문자열입니다."    

# === Streamlit 앱 ===
st.title("📧 EML 본문 조회")

# 데이터베이스 연결 확인
try:
    conn = sqlite3.connect(EML_DB_NAME)
    conn.close()
except Exception as e:
    st.error(f"❌ 데이터베이스 연결 실패: {str(e)}")
    st.stop()

# 레코드 조회
records = get_all_eml_records()

if not records:
    st.info("조회된 레코드가 없습니다.")
else:
    # 레코드 선택
    record_options = [f"ID {record[0]} - {record[1]}" for record in records]
    selected_index = st.selectbox("조회할 레코드를 선택하세요:", range(len(records)), format_func=lambda x: record_options[x])
    
    # 선택된 레코드의 본문 표시
    if selected_index is not None:
        selected_record = records[selected_index]
        body_text = selected_record[3]
        
        # 본문 내용 표시
        st.subheader("본문 내용")
        with st.expander("본문 내용 보기", expanded=True):
            st.markdown(f"```\n{body_text}\n```")
        
        # 샘플 장애보고서 내용 표시
        st.subheader("샘플 장애보고서")
        sample_content = read_sample_document("data/docx/iap-report-sample1.docx")
        with st.expander("샘플 장애보고서 내용 보기", expanded=True):
            st.markdown(f"```\n{sample_content}\n```")

        # LLM 생성 보고서 섹션 추가
        st.subheader("🤖 LLM 생성 보고서")

        with st.spinner("LLM이 보고서를 생성하고 있습니다..."):
            llm_report = generate_llm_report(body_text, sample_content)
        
        with st.expander("LLM 생성 보고서 보기", expanded=True):
            st.markdown(llm_report)

        st.markdown("---")
        if st.button("📄 Word 파일 생성 및 업로드"):
            word_file = generate_word_from_body(body_text, "generated_report.docx")
            blob_name = f"iap-report-{selected_record[0]}.docx"
            blob_url = upload_to_azure_word_blob(word_file, blob_name)

            if blob_url:
                st.success("✅ Word 파일이 생성되어 Azure에 업로드되었습니다.")
                st.markdown(f"[📥 파일 다운로드]({blob_url})", unsafe_allow_html=True)